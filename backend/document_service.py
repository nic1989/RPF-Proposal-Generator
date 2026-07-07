"""Extract text from PDF/DOCX uploads and export proposals to DOCX/PDF."""
from __future__ import annotations
import io
from typing import List

from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", errors="ignore")
    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    out = []
    for page in reader.pages:
        try:
            out.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(out).strip()


def _extract_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: List[str] = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def build_proposal_docx(title: str, exec_summary: str, sections: list, company: str | None = None) -> bytes:
    doc = Document()

    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    if company:
        p = doc.add_paragraph(f"Prepared by: {company}")
        p.runs[0].font.size = Pt(11)

    if exec_summary:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(exec_summary)

    for sec in sections:
        doc.add_heading(sec.get("heading", "Section"), level=1)
        for para in (sec.get("content", "") or "").split("\n"):
            if para.strip():
                doc.add_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_proposal_pdf(title: str, exec_summary: str, sections: list, company: str | None = None) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER, leftMargin=54, rightMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("T", parent=styles["Title"], alignment=TA_LEFT, fontSize=22, textColor="#09090B")
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=15, textColor="#09090B", spaceAfter=8)
    body = ParagraphStyle("B", parent=styles["BodyText"], fontSize=11, leading=16, textColor="#27272A")
    meta = ParagraphStyle("M", parent=styles["BodyText"], fontSize=10, textColor="#52525B")

    story = [Paragraph(title, title_style), Spacer(1, 8)]
    if company:
        story.append(Paragraph(f"Prepared by: {company}", meta))
        story.append(Spacer(1, 12))
    if exec_summary:
        story.append(Paragraph("Executive Summary", h1))
        for para in exec_summary.split("\n"):
            if para.strip():
                story.append(Paragraph(para.replace("&", "&amp;"), body))
                story.append(Spacer(1, 6))

    for sec in sections:
        story.append(Spacer(1, 8))
        story.append(Paragraph(sec.get("heading", "Section"), h1))
        for para in (sec.get("content", "") or "").split("\n"):
            if para.strip():
                story.append(Paragraph(para.replace("&", "&amp;"), body))
                story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()
